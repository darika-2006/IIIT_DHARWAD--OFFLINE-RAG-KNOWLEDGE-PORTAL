from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from typing import List
from sentence_transformers import SentenceTransformer
import pypdfium2 as pdfium
import faiss  #faisss - cpu for library installation 
import psycopg2
import os
from docx import Document #docx-python 
from passlib.context import CryptContext
from pydantic import BaseModel
from fastapi import HTTPException
from fastapi.middleware.cors import CORSMiddleware



#for bcrypt :
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Server connection (APP) 
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# DB 
conn = psycopg2.connect(
    dbname="RAGY",          # your database name
    user="postgres",      # postgres role
    password="Bala@2007",
    host="localhost",
    port="5432"
)
cur = conn.cursor()

#Model :
model = SentenceTransformer("all-MiniLM-L6-v2")

#TEXT EXTRACTION 
def extract_text_from_pdf(pdf_path, max_words=250, overlap=50):
    pdf = pdfium.PdfDocument(pdf_path)
    chunks = []
    for page_num in range(len(pdf)):
        page = pdf.get_page(page_num)
        text_page = page.get_textpage()
        page_text = text_page.get_text_range()
        text_page.close()
        page.close()

        words = page_text.split()
        start = 0

        while start < len(words):
            end = start + max_words
            chunk_text = " ".join(words[start:end])

            chunks.append({
                "page_number": page_num + 1,
                "text": chunk_text
            })

            start = end - overlap
            if overlap >= max_words:
                break
    pdf.close()
    return chunks

def extract_text_from_docx(docx_path, max_words=250, overlap=50):
    doc = Document(docx_path)
    full_text = " ".join([p.text for p in doc.paragraphs])

    words = full_text.split()
    chunks = []
    start = 0
    page_number = 1  # logical page

    while start < len(words):
        end = start + max_words
        chunk_text = " ".join(words[start:end])

        chunks.append({
            "page_number": page_number,
            "text": chunk_text
        })

        start = end - overlap
        page_number += 1

    return chunks

def extract_text_from_txtFile(txt_path,max_words=250,overlap=50):
    chunks=[]
    file_path = txt_path
    try:
        with open(file_path, 'r') as file:
            words= file.readlines()
    except FileNotFoundError:
        HTTPException(status_code=400,detail="File Doesn't Exist !")
    
    start=0
    pg_no=1
    while(start<len(words)):
        end=start+max_words
        chunk_text=" ".join(words[start:end])
        chunks.append({
            "page_number":pg_no,
            "text":chunk_text
        })
        start=end-overlap
        pg_no+=1
    return chunks

# CHUNK + DB 
def chunk_list(file_paths, allowed_roles, uploaded_by="admin"):
    chunks = []
    for path in file_paths:
        file_name = os.path.basename(path)
        sql = """
        INSERT INTO documents (doc_name, doc_path, allowed_roles, uploaded_by)
        VALUES (%s, %s, %s, %s)
        RETURNING doc_id
        """
        cur.execute(sql, (file_name, path, allowed_roles, uploaded_by))
        conn.commit()

        doc_id = cur.fetchone()[0]

        if path.endswith(".pdf"):
            extracted = extract_text_from_pdf(path)
        elif path.endswith(".docx"):
            extracted = extract_text_from_docx(path)
        elif path.endswith(".txt"):
            extracted=extract_text_from_txtFile(path)
        else:
            raise HTTPException(status_code=400,detail="Unsupported file type")
        for chunk in extracted:
            chunk["doc_id"] = doc_id
            chunk["allowed_roles"] = allowed_roles
            chunks.append(chunk)
    return chunks


def build_metadata(chunks):
    cur.execute("SELECT COALESCE(MAX(chunk_id), -1) FROM chunks_metadata")
    base_chunk_id = cur.fetchone()[0] + 1

    sql = """
    INSERT INTO chunks_metadata
    (chunk_id, doc_id, page_number, chunk_text, allowed_roles, is_active)
    VALUES (%s, %s, %s, %s, %s, TRUE)
    """

    for i, chunk in enumerate(chunks):
        cur.execute(sql,(
            base_chunk_id + i,
            chunk["doc_id"],
            chunk["page_number"],
            chunk["text"],
            chunk["allowed_roles"]
        ))
    conn.commit()

# UPLOAD API 

@app.post("/upload")
def upload_documents(
    files: List[UploadFile] = File(...),
    allowed_roles: str = Form(...),
    uploaded_by: str = Form("admin")
):
    saved_paths = []
    for file in files:
        if not file.filename.endswith((".pdf", ".docx",".txt")):
            raise HTTPException(status_code=400, detail="Only PDF and DOCX allowed")
        path = os.path.join(UPLOAD_DIR, file.filename)
        with open(path, "wb") as f:
            f.write(file.file.read())
        saved_paths.append(path)
    chunks = chunk_list(saved_paths, allowed_roles, uploaded_by)
    if not chunks:
        raise HTTPException(status_code=400, detail="No text extracted")

    texts = [c["text"] for c in chunks]
    embeddings = model.encode(texts, normalize_embeddings=True)

    dim = embeddings.shape[1]
    if os.path.exists("halx_index.faiss"):
        index = faiss.read_index("halx_index.faiss")
    else:
        index = faiss.IndexFlatL2(dim)

    base_chunk_id = index.ntotal
    index.add(embeddings)

    build_metadata(chunks)
    faiss.write_index(index, "halx_index.faiss")

    return {
        "message": "Documents indexed successfully",
        "documents_uploaded": len(files),
        "chunks_created": len(chunks)
    }

# Delete 
def _delete_metadata(doc_id: int):
    sql_doc = "UPDATE documents SET is_active = FALSE WHERE doc_id = %s"
    sql_chunks = "UPDATE chunks_metadata SET is_active = FALSE WHERE doc_id = %s"

    cur.execute(sql_doc, (doc_id,))
    cur.execute(sql_chunks, (doc_id,))
    conn.commit()


@app.post("/delete-document")
def delete_document(doc_id: int):
    cur.execute("SELECT doc_id FROM documents WHERE doc_id = %s", (doc_id,))
    result = cur.fetchone()

    if not result:
        raise HTTPException(status_code=404, detail="Document not found")

    _delete_metadata(doc_id)

    return {
        "message": "Document deleted successfully",
        "doc_id": doc_id
    }

class SignupRequest(BaseModel):
    user_id: str
    username: str
    password: str
    domain_role: str  # admin / user


class LoginRequest(BaseModel):
    username: str
    password: str

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

@app.post("/signup")
def signup(data: SignupRequest):
    # Check if username exists
    cur.execute(
        "SELECT username FROM login_credentials WHERE username = %s",
        (data.username,)
    )
    if cur.fetchone():
        raise HTTPException(status_code=400, detail="Username already exists")

    hashed_pw = hash_password(data.password)

    try:
        cur.execute(
            """
            INSERT INTO login_credentials (user_id, username, password_hash, domain_role)
            VALUES (%s, %s, %s, %s)
            """,
            (data.user_id, data.username, hashed_pw, data.domain_role)
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail="Signup failed")

    return {
        "message": "Signup successful",
        "username": data.username,
        "role": data.domain_role
    }

@app.post("/login")
def login(data: LoginRequest):
    cur.execute(
        """
        SELECT user_id, username, password_hash, domain_role
        FROM login_credentials
        WHERE username = %s
        """,
        (data.username,)
    )
    user = cur.fetchone()

    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")

    user_id, username, password_hash, role = user

    if not verify_password(data.password, password_hash):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    return {
        "message": "Login successful",
        "user_id": user_id,
        "username": username,
        "role": role
    }

